# modules/generar_reporte.py
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import io

def crear_docx(output_path, indices, gdf, datos_suelo, curvas_img_path, recomendaciones):
    doc = Document()

    # Título
    title = doc.add_heading('Informe Agronómico de Lote', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sección 1: Datos generales
    doc.add_heading('1. Datos Generales', level=2)
    area = gdf.to_crs("EPSG:32719").area.sum() / 10000
    doc.add_paragraph(f"Área total del lote: {area:.2f} hectáreas.")
    # Podrías agregar más info del GeoJSON (ubicación aproximada)

    # Sección 2: Índices espectrales
    doc.add_heading('2. Índices de Vegetación y Agua', level=2)
    p = doc.add_paragraph()
    p.add_run(f"NDVI promedio: {indices['ndvi_mean']:.3f}\n").bold = True
    p.add_run("Interpretación: ")
    if indices['ndvi_mean'] > 0.6:
        p.add_run("Vegetación vigorosa y saludable.\n")
    elif indices['ndvi_mean'] > 0.3:
        p.add_run("Vegetación moderada, posible estrés o baja densidad.\n")
    else:
        p.add_run("Baja cobertura vegetal o suelo desnudo.\n")

    p.add_run(f"NDWI promedio: {indices['ndwi_mean']:.3f}\n").bold = True
    if indices['ndwi_mean'] > 0:
        p.add_run("Presencia de humedad o agua en la vegetación.\n")
    else:
        p.add_run("Condiciones secas.\n")

    # Sección 3: Datos de suelo (si existen)
    if datos_suelo:
        doc.add_heading('3. Análisis de Suelo', level=2)
        tabla = doc.add_table(rows=1, cols=2)
        tabla.style = 'Light Grid Accent 1'
        hdr_cells = tabla.rows[0].cells
        hdr_cells[0].text = 'Parámetro'
        hdr_cells[1].text = 'Valor promedio'
        for param, valor in datos_suelo.items():
            row_cells = tabla.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = f"{valor:.2f}"
    else:
        doc.add_heading('3. Datos de Suelo', level=2)
        doc.add_paragraph("No se proporcionaron datos de suelo.")

    # Sección 4: Topografía (si hay imagen de curvas)
    if curvas_img_path:
        doc.add_heading('4. Topografía y Curvas de Nivel', level=2)
        doc.add_picture(curvas_img_path, width=Inches(6))
        doc.add_paragraph("Figura 1: Curvas de nivel dentro del lote.")
    else:
        doc.add_heading('4. Topografía', level=2)
        doc.add_paragraph("No se proporcionó modelo de elevación.")

    # Sección 5: Recomendaciones de IA
    doc.add_heading('5. Análisis y Recomendaciones', level=2)
    doc.add_paragraph(recomendaciones)

    # Guardar
    doc.save(output_path)
